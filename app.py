import os
import io
import json
import re
import streamlit as st
import pandas as pd
from openai import OpenAI
from docx import Document

# ============================
# App Config
# ============================
st.set_page_config(page_title="Safety Topic Analyzer", layout="wide")
st.title("Safety Topic Analyzer")

SYSTEM_STYLE = """
You are acting as an EU QPPV-level safety physician.
You are experienced in signal validation and signal assessment, and your writing may be reviewed by EMA, MHRA, or FDA.

Writing rules:
- Inspection-ready and conservative.
- Evidence-based only; avoid speculation.
- Neutral tone; no persuasive language.
- If data is insufficient, explicitly state limitations.
""".strip()

# ============================
# OpenAI Client
# ============================
api_key = None
try:
    api_key = st.secrets.get("OPENAI_API_KEY")
except Exception:
    api_key = None

if not api_key:
    api_key = os.getenv("OPENAI_API_KEY")

if not api_key:
    st.error("OpenAI API key not found. Add OPENAI_API_KEY in Streamlit Cloud → Settings → Secrets.")
    st.stop()

client = OpenAI(api_key=api_key)

def call_openai(user_prompt: str, model: str = "gpt-5") -> str:
    r = client.chat.completions.create(
        model=model,
        messages=[
            {"role": "system", "content": SYSTEM_STYLE},
            {"role": "user", "content": user_prompt},
        ],
    )
    return r.choices[0].message.content

# ============================
# Helpers
# ============================
def find_col(df, candidates):
    cols = list(df.columns)
    cols_l = [str(c).lower().strip() for c in cols]
    for cand in candidates:
        cand_l = cand.lower()
        for i, c in enumerate(cols_l):
            if c == cand_l:
                return cols[i]
        for i, c in enumerate(cols_l):
            if cand_l in c:
                return cols[i]
    return None

def yes_count(series):
    s = series.astype(str).str.strip().str.lower()
    return int(s.isin(["y", "yes", "true", "1"]).sum())

def top_values(series, n=5):
    return series.dropna().astype(str).value_counts().head(n).to_dict()

def json_or_none(s: str):
    try:
        return json.loads(s)
    except Exception:
        return None

def must_have_keys(d, keys):
    return isinstance(d, dict) and all(k in d for k in keys)

def rating_norm(s: str) -> str:
    if not s:
        return ""
    s = str(s).strip().lower()
    if "not" in s and "assess" in s:
        return "Not assessable"
    if "weak" in s:
        return "Weak"
    if "moderate" in s:
        return "Moderate"
    if "strong" in s:
        return "Strong"
    if "low" in s:
        return "Low"
    if "high" in s:
        return "High"
    return s.title()

def ensure_single_paragraph(text: str) -> str:
    lines = [ln.strip() for ln in str(text).splitlines() if ln.strip()]
    cleaned = []
    for ln in lines:
        if ln.startswith(("-", "•", "*")):
            ln = ln.lstrip("-•* ").strip()
        cleaned.append(ln)
    para = " ".join(cleaned).strip()
    while "  " in para:
        para = para.replace("  ", " ")
    return para

def final_signal_outcome(causality: str) -> str:
    c = (causality or "").strip().lower()
    if c in ["supported", "possible"]:
        return "Signal Confirmed"
    return "Signal Refuted"

# ============================
# Upload Excel
# ============================
uploaded_file = st.file_uploader("Upload Excel Line Listing", type=["xlsx"])
if uploaded_file is None:
    st.info("Upload an Excel line listing to begin.")
    st.stop()

df = pd.read_excel(uploaded_file)
st.subheader("Preview of Uploaded Data")
st.dataframe(df.head(20))

# ============================
# Column detection
# ============================
PT_CANDS = ["event pt", "preferred term", "meddra pt", "reaction pt", "pt"]
SERIOUS_CANDS = ["serious case flag", "serious", "event seriousness", "seriousness"]
FATAL_CANDS = ["death flag", "fatal case flag", "fatal"]
LISTED_CANDS = ["listedness", "event listedness", "expected"]
DECH_CANDS = ["dechallenge", "dechallenge results"]
RECH_CANDS = ["rechallenge", "rechallenge results"]
ONSET_CANDS = ["onset latency", "time to onset", "event start date"]
COUNTRY_CANDS = ["country"]
NARR_CANDS = ["narrative"]
DRUG_CANDS = ["suspect product name", "suspect drug", "drug", "product", "generic name"]
CASE_CANDS = ["case number", "case id", "icsr", "report id", "safety report id"]

pt_col = find_col(df, PT_CANDS)
ser_col = find_col(df, SERIOUS_CANDS)
fat_col = find_col(df, FATAL_CANDS)
list_col = find_col(df, LISTED_CANDS)
dech_col = find_col(df, DECH_CANDS)
rech_col = find_col(df, RECH_CANDS)
onset_col = find_col(df, ONSET_CANDS)
country_col = find_col(df, COUNTRY_CANDS)
narr_col = find_col(df, NARR_CANDS)
drug_col = find_col(df, DRUG_CANDS)
case_col = find_col(df, CASE_CANDS)

if not pt_col:
    st.error("Event PT column not detected. Rename column to include 'Event PT' or 'Preferred Term'.")
    st.stop()

topic_table = (
    df.groupby(pt_col, dropna=True)
    .size()
    .reset_index(name="row_count")
    .sort_values("row_count", ascending=False)
)
st.subheader("Auto-grouped Topics (by Event PT)")
st.dataframe(topic_table.head(50))

# ============================
# Sidebar: Product context
# ============================
st.sidebar.header("Product Context (stabilizes Background)")
product_name = st.sidebar.text_input("Suspect Product (brand or substance)", value="")
generic_name = st.sidebar.text_input("Generic name (optional)", value="")
ther_class = st.sidebar.text_input("Therapeutic class (optional)", value="")

def product_context_string():
    parts = []
    if product_name.strip():
        parts.append(f"Suspect product: {product_name.strip()}")
    if generic_name.strip():
        parts.append(f"Generic name: {generic_name.strip()}")
    if ther_class.strip():
        parts.append(f"Therapeutic class: {ther_class.strip()}")
    if not parts:
        parts.append("Suspect product: Not provided (keep background high-level and cautious).")
    return "\n".join(parts)

# ============================
# Evidence builder
# ============================
def build_evidence(subset, pt_value):
    evidence = {
        "event_pt": str(pt_value),
        "row_count": int(len(subset)),
        "unique_case_count": int(subset[case_col].nunique()) if case_col else None,
        "serious_yes": yes_count(subset[ser_col]) if ser_col else None,
        "fatal_yes": yes_count(subset[fat_col]) if fat_col else None,
        "listedness_top": top_values(subset[list_col]) if list_col else None,
        "dechallenge_top": top_values(subset[dech_col]) if dech_col else None,
        "rechallenge_top": top_values(subset[rech_col]) if rech_col else None,
        "onset_examples": subset[onset_col].dropna().astype(str).head(5).tolist() if onset_col else [],
        "countries_top": top_values(subset[country_col]) if country_col else None,
        "narrative_snippets": subset[narr_col].dropna().astype(str).head(3).tolist() if narr_col else [],
    }
    if drug_col:
        evidence["suspect_products_top"] = top_values(subset[drug_col])
    return evidence

# ============================
# Modes
# ============================
st.subheader("Choose Output Mode")
mode = st.radio(
    "Mode",
    [
        "Signal Assessment (with WoE table + binary outcome)",
        "Bulk Trending (Monthly Scan)",
        "PBRER Summary (Aggregate Synopsis)",
    ],
)

# ==================================================
# MODE 1: Signal Assessment
# ==================================================
if mode.startswith("Signal Assessment"):
    st.subheader("Signal Assessment (Structured + WoE + Binary Outcome)")

    pt_choice = st.selectbox("Select Event PT", topic_table[pt_col].tolist())

    if st.button("Generate Signal Assessment + Word"):
        subset = df[df[pt_col] == pt_choice]
        evidence = build_evidence(subset, pt_choice)

        # ---- Step 1: WoE scoring JSON ----
        score_keys = [
            "Strength",
            "Consistency",
            "Temporality",
            "Plausibility",
            "Confounding",
            "Overall Evidence",
            "Rationale (brief)",
            "Causality Conclusion",
        ]

        score_prompt = f"""
Return ONLY valid JSON with EXACT keys:

{{
  "Strength": "Weak/Moderate/Strong",
  "Consistency": "Weak/Moderate/Strong",
  "Temporality": "Weak/Moderate/Strong/Not assessable",
  "Plausibility": "Weak/Moderate/Strong",
  "Confounding": "Low/Moderate/High",
  "Overall Evidence": "Weak/Moderate/Strong",
  "Rationale (brief)": "2-3 sentences, evidence-based, inspection-ready",
  "Causality Conclusion": "Supported/Possible/Insufficient Evidence/Unlikely"
}}

Use only provided evidence. No speculation.

PRODUCT CONTEXT:
{product_context_string()}

EVIDENCE:
{evidence}
"""
        raw_scores = call_openai(score_prompt)
        scores = json_or_none(raw_scores)
        if scores is None or not must_have_keys(scores, score_keys):
            repair = call_openai("Convert into ONLY valid JSON with the exact keys. Output JSON only.\n\n" + raw_scores)
            scores = json_or_none(repair)

        if scores is None or not must_have_keys(scores, score_keys):
            st.error("Could not parse scoring JSON. Raw output below:")
            st.code(raw_scores)
            st.stop()

        # Normalize ratings + binary outcome
        for k in ["Strength", "Consistency", "Temporality", "Plausibility", "Confounding", "Overall Evidence"]:
            scores[k] = rating_norm(scores.get(k, ""))

        causality = str(scores.get("Causality Conclusion", "")).strip()
        signal_outcome = final_signal_outcome(causality)

        # ---- Step 2: Structured narrative JSON (NO PBRER summary here) ----
        section_keys = [
            "Background of Drug-Event Combination",
            "Case Synopsis",
            "Bradford Hill Assessment",
            "Regulatory Implications",
            "Causality Conclusion",
            "Final Signal Outcome",
        ]

        structured_prompt = f"""
Return ONLY valid JSON (no markdown) with EXACT keys:

{{
  "Background of Drug-Event Combination": "string",
  "Case Synopsis": "string",
  "Bradford Hill Assessment": "string",
  "Regulatory Implications": "string",
  "Causality Conclusion": "string",
  "Final Signal Outcome": "string"
}}

RULES:
- Output ALL keys in EXACT order above.
- Background: product-level only, <=120 words, single paragraph, no case counts.
- Regulatory Implications: NO browsing; do not invent safety communications.
  If uncertain, write exactly:
  "No widely recognized regulatory action specific to this drug-event combination."
- Use these EXACT values:
  Causality Conclusion: "{causality}"
  Final Signal Outcome: "{signal_outcome}"
- Keep sections as short paragraphs. Avoid bullet-only sections.

PRODUCT CONTEXT:
{product_context_string()}

Woe Scores (context only):
{scores}

REPORTING INTERVAL EVIDENCE:
{evidence}
"""
        raw_struct = call_openai(structured_prompt)
        data = json_or_none(raw_struct)
        if data is None or not must_have_keys(data, section_keys):
            repair = call_openai("Convert into ONLY valid JSON with the exact keys. Output JSON only.\n\n" + raw_struct)
            data = json_or_none(repair)

        if data is None or not must_have_keys(data, section_keys):
            st.error("Could not parse structured JSON. Raw output below:")
            st.code(raw_struct)
            st.stop()

        data["Background of Drug-Event Combination"] = ensure_single_paragraph(data.get("Background of Drug-Event Combination", ""))

        # ---- Display in app ----
        st.markdown("### Signal Weight-of-Evidence (WoE) Summary")
        woe_df = pd.DataFrame(
            [
                ["Strength", scores.get("Strength", "")],
                ["Consistency", scores.get("Consistency", "")],
                ["Temporality", scores.get("Temporality", "")],
                ["Plausibility", scores.get("Plausibility", "")],
                ["Confounding", scores.get("Confounding", "")],
                ["Overall Evidence", scores.get("Overall Evidence", "")],
            ],
            columns=["Criterion", "Rating"],
        )
        st.table(woe_df)

        st.markdown("### Rationale (brief)")
        st.write(str(scores.get("Rationale (brief)", "")).strip())

        st.markdown("### Final Outputs")
        st.write({"Causality Conclusion": causality, "Final Signal Outcome": signal_outcome})

        st.subheader(f"Signal Assessment – {pt_choice}")
        for k in section_keys:
            st.markdown(f"### {k}")
            st.write(str(data.get(k, "")).strip())

        # ---- Word export ----
        doc = Document()
        doc.add_heading("Signal Assessment Report", level=1)
        doc.add_paragraph(f"Event PT: {pt_choice}")
        doc.add_paragraph("")

        doc.add_heading("Signal Weight-of-Evidence (WoE) Summary", level=2)
        t = doc.add_table(rows=1, cols=2)
        t.rows[0].cells[0].text = "Criterion"
        t.rows[0].cells[1].text = "Rating"
        for row in [
            ("Strength", scores.get("Strength", "")),
            ("Consistency", scores.get("Consistency", "")),
            ("Temporality", scores.get("Temporality", "")),
            ("Plausibility", scores.get("Plausibility", "")),
            ("Confounding", scores.get("Confounding", "")),
            ("Overall Evidence", scores.get("Overall Evidence", "")),
        ]:
            r = t.add_row().cells
            r[0].text = row[0]
            r[1].text = str(row[1])

        doc.add_paragraph("")
        doc.add_heading("Rationale (brief)", level=2)
        doc.add_paragraph(str(scores.get("Rationale (brief)", "")).strip())

        doc.add_heading("Final Outputs", level=2)
        doc.add_paragraph(f"Causality Conclusion: {causality}")
        doc.add_paragraph(f"Final Signal Outcome: {signal_outcome}")

        doc.add_page_break()
        doc.add_heading("Structured Assessment", level=2)
        for k in section_keys:
            doc.add_heading(k, level=3)
            doc.add_paragraph(str(data.get(k, "")).strip())

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)

        st.download_button(
            "Download Signal Assessment (Word)",
            data=buf,
            file_name=f"Signal_Assessment_{pt_choice}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

# ==================================================
# MODE 2: Bulk Trending (unchanged)
# ==================================================
elif mode.startswith("Bulk Trending"):
    st.subheader("Bulk Trending (Monthly Scan)")
    TOP_N = st.slider("Number of PTs to Analyze", 3, 30, 10)

    if st.button("Analyze Top PTs"):
        top_pts = topic_table.head(TOP_N)[pt_col].tolist()

        for pt in top_pts:
            subset = df[df[pt_col] == pt]
            evidence = build_evidence(subset, pt)

            prompt = f"""
Provide:
1) Event PT
2) Safety topic decision (Include / Continue Monitoring / Close / Escalate)
3) Key evidence (max 4 bullets)
4) Causality conclusion (Supported / Possible / Insufficient Evidence / Unlikely)
5) Recommended action

Evidence:
{evidence}
"""
            out = call_openai(prompt)
            st.subheader(f"Trending – {pt}")
            st.write(out)
            st.markdown("---")

# ==================================================
# MODE 3: PBRER Summary (optional)
# ==================================================
else:
    st.subheader("PBRER Summary (Aggregate Synopsis)")

    pt_choice = st.selectbox("Select Event PT", topic_table[pt_col].tolist())

    if st.button("Generate PBRER Summary + Word"):
        subset = df[df[pt_col] == pt_choice]
        evidence = build_evidence(subset, pt_choice)

        prompt = f"""
Write a PBRER-ready aggregate synopsis for the reporting interval.

Return ONLY valid JSON with EXACT keys:
{{
  "PBRER Summary": "string"
}}

RULES for "PBRER Summary":
- ONE continuous paragraph (no bullets, no line breaks, no labels)
- 140–190 words
- Include: case counts, seriousness, fatalities, geography, listedness (if available), key clinical pattern(s), confounders/alternative etiologies, and an overall benefit–risk position sentence at the end.
- Do NOT mention Bradford Hill explicitly.
- Do NOT invent regulatory actions.
- Use only evidence provided.

PRODUCT CONTEXT:
{product_context_string()}

EVIDENCE:
{evidence}
"""
        raw = call_openai(prompt)
        data = json_or_none(raw)
        if data is None or "PBRER Summary" not in data:
            repair = call_openai("Convert into ONLY valid JSON with the exact key. Output JSON only.\n\n" + raw)
            data = json_or_none(repair)

        if data is None or "PBRER Summary" not in data:
            st.error("Could not parse PBRER Summary JSON. Raw output below:")
            st.code(raw)
            st.stop()

        summary = ensure_single_paragraph(data["PBRER Summary"])
        st.subheader(f"PBRER Summary – {pt_choice}")
        st.write(summary)

        doc = Document()
        doc.add_heading("PBRER Summary (Aggregate Synopsis)", level=1)
        doc.add_paragraph(f"Event PT: {pt_choice}")
        doc.add_paragraph("")
        doc.add_paragraph(summary)

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        st.download_button(
            "Download PBRER Summary (Word)",
            data=buf,
            file_name=f"PBRER_Summary_{pt_choice}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
